from __future__ import annotations

import json
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "documentos" / "Paginas_da_Vida_Revisao_Pastoral.docx"
NODE = Path(
    r"C:\Users\Lucas\.cache\codex-runtimes\codex-primary-runtime\dependencies\node\bin\node.exe"
)

NAVY = RGBColor(27, 41, 56)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GOLD = RGBColor(154, 123, 69)
GRAY = RGBColor(95, 99, 104)
LIGHT_GRAY = RGBColor(225, 228, 232)
BLACK = RGBColor(28, 28, 28)


def load_stories() -> dict:
    source = ROOT / "js" / "historias.js"
    js = (
        "global.window={};"
        f"require({json.dumps(str(source))});"
        "process.stdout.write(JSON.stringify(window.HISTORIAS));"
    )
    result = subprocess.run(
        [str(NODE), "-e", js],
        check=True,
        capture_output=True,
        text=True,
        encoding="utf-8",
    )
    return json.loads(result.stdout)


def set_font(run, size=None, bold=None, italic=None, color=BLACK, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = color


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    set_font(run, 9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = paragraph.add_run()
    field_run._r.extend([begin, instr, separate, value, end])
    set_font(field_run, 9, color=GRAY)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 18, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 9, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("PÁGINAS DA VIDA  |  JSO")
    set_font(run, 8.5, bold=True, color=GOLD)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_rule(paragraph, color="9A7B45", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_cover(doc: Document):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(90)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_font(kicker.add_run("EXPERIÊNCIA INTERATIVA • JSO"), 10, bold=True, color=GOLD)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_font(title.add_run("PÁGINAS DA VIDA"), 30, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    set_font(subtitle.add_run("Protagonista ou coadjuvante?"), 16, italic=True, color=DARK_BLUE)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.left_indent = Inches(1.4)
    line.paragraph_format.right_indent = Inches(1.4)
    add_rule(line)

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.paragraph_format.space_before = Pt(18)
    desc.paragraph_format.space_after = Pt(110)
    set_font(
        desc.add_run(
            "Documento para revisão pastoral dos 30 capítulos\n"
            "Experiência • Revelação • Conclusão"
        ),
        12,
        color=GRAY,
    )

    prepared = doc.add_paragraph()
    prepared.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(prepared.add_run("Preparado para o Pastor Renato"), 11, bold=True, color=NAVY)

    ministry = doc.add_paragraph()
    ministry.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(ministry.add_run("Juventude Superando Obstáculos"), 10, color=GRAY)
    doc.add_page_break()


def add_intro(doc: Document):
    doc.add_heading("Orientação para revisão", level=1)
    paragraphs = [
        (
            "Este documento reúne os 30 capítulos da experiência Páginas da Vida. "
            "Cada história representa apenas uma página de um roteiro maior e reforça "
            "a mensagem central: nós enxergamos o capítulo presente, enquanto Deus conhece a história completa."
        ),
        (
            "Cada capítulo possui três momentos: Experiência, para gerar identificação; "
            "Revelação, para interpretar a primeira leitura à luz da soberania de Deus; "
            "e Conclusão, para conduzir a uma aplicação de fé acompanhada de uma referência bíblica."
        ),
        (
            "Sugestões de revisão: clareza pastoral, fidelidade bíblica, sensibilidade emocional, "
            "linguagem adequada aos jovens adultos e coerência com o tema do culto."
        ),
    ]
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(16)
    set_font(note.add_run("Status geral:  [  ] Aprovado   [  ] Ajustes necessários"), 10.5, bold=True, color=NAVY)


def add_index(doc: Document, stories: dict):
    doc.add_heading("Sumário dos capítulos", level=1)
    sorted_items = sorted(stories.items(), key=lambda item: int(item[0]))
    for number, story in sorted_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        set_font(p.add_run(f"Capítulo {number}"), 10, bold=True, color=GOLD)
        set_font(p.add_run(f"  —  {story['titulo']}"), 10, color=BLACK)
    doc.add_page_break()


def add_story_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_together = True
    parts = text.split("\n")
    for index, part in enumerate(parts):
        set_font(p.add_run(part), 10.5, color=BLACK)
        if index < len(parts) - 1:
            p.add_run().add_break()


def add_moment(doc: Document, label: str, moment: dict, page_number: int):
    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 2"]
    set_font(heading.add_run(label.upper()), 13, bold=True, color=GOLD)

    for text in moment["paragrafos"]:
        add_story_paragraph(doc, text)

    if moment.get("citacao"):
        quote = doc.add_paragraph()
        quote.paragraph_format.left_indent = Inches(0.35)
        quote.paragraph_format.right_indent = Inches(0.25)
        quote.paragraph_format.space_before = Pt(8)
        quote.paragraph_format.space_after = Pt(2)
        set_font(quote.add_run(moment["citacao"]), 10.5, italic=True, color=DARK_BLUE)

        ref = doc.add_paragraph()
        ref.paragraph_format.left_indent = Inches(0.35)
        ref.paragraph_format.space_after = Pt(8)
        set_font(ref.add_run(moment["referencia"]), 9.5, bold=True, color=GOLD)

    final = doc.add_paragraph()
    final.paragraph_format.space_before = Pt(5)
    final.paragraph_format.space_after = Pt(3)
    set_font(final.add_run(moment["fraseFinal"]), 10, bold=True, color=NAVY)

    page = doc.add_paragraph()
    page.paragraph_format.space_after = Pt(10)
    set_font(page.add_run(f"Página {page_number} de ?"), 9, italic=True, color=GRAY)

    if page_number == 3:
        epilogue = doc.add_paragraph()
        epilogue.paragraph_format.space_before = Pt(4)
        epilogue.paragraph_format.space_after = Pt(10)
        set_font(
            epilogue.add_run("O restante da história continua sendo escrito pelo Diretor."),
            10,
            bold=True,
            color=NAVY,
        )


def add_chapter(doc: Document, number: str, story: dict, first: bool):
    if not first:
        doc.add_page_break()

    chapter = doc.add_paragraph()
    chapter.paragraph_format.space_after = Pt(2)
    chapter.paragraph_format.keep_with_next = True
    set_font(chapter.add_run(f"CAPÍTULO {number}"), 10, bold=True, color=GOLD)

    title = doc.add_heading(story["titulo"], level=1)
    title.paragraph_format.space_before = Pt(0)

    moments = story["momentos"]
    add_moment(doc, "Momento 1 — Experiência", moments["experiencia"], 1)
    add_moment(doc, "Momento 2 — Revelação", moments["revelacao"], 2)
    add_moment(doc, "Momento 3 — Conclusão", moments["conclusao"], 3)

    review = doc.add_paragraph()
    review.paragraph_format.space_before = Pt(10)
    review.paragraph_format.space_after = Pt(4)
    set_font(review.add_run("VALIDAÇÃO PASTORAL"), 9.5, bold=True, color=GOLD)

    status = doc.add_paragraph()
    status.paragraph_format.space_after = Pt(6)
    set_font(status.add_run("[  ] Aprovado   [  ] Ajustar   [  ] Revisar referência bíblica"), 9.5, color=BLACK)

    obs = doc.add_paragraph()
    obs.paragraph_format.space_after = Pt(0)
    set_font(obs.add_run("Observações: ______________________________________________________________"), 9.5, color=GRAY)
    obs2 = doc.add_paragraph()
    obs2.paragraph_format.space_after = Pt(0)
    set_font(obs2.add_run("__________________________________________________________________________"), 9.5, color=GRAY)


def main():
    stories = load_stories()
    if len(stories) != 30:
        raise RuntimeError(f"Esperados 30 capítulos; encontrados {len(stories)}.")

    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_intro(doc)
    add_index(doc, stories)

    items = sorted(stories.items(), key=lambda item: int(item[0]))
    for index, (number, story) in enumerate(items):
        add_chapter(doc, number, story, first=index == 0)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = "Páginas da Vida — Revisão Pastoral"
    doc.core_properties.subject = "30 capítulos da experiência interativa da JSO"
    doc.core_properties.author = "Juventude Superando Obstáculos"
    doc.core_properties.keywords = "JSO, Páginas da Vida, revisão pastoral"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
